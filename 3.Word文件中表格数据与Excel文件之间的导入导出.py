from random import choice
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import load_workbook, Workbook

def xlsx2docx(fn):
    # 打开Excel文件，如果有公式的话，读取公式计算结果
    wb = load_workbook(fn, data_only=True)
    # 创建空白Word文件
    document = Document()
    
    # 查看所有可用的表格样式
    table_styles = [style for style in document.styles if style.type==WD_STYLE_TYPE.TABLE]
    
    # 遍历Excel文件中所有的worksheet
    for ws in wb.worksheets:
        rows = list(ws.rows)
        # 增加段落，也就是表格的名称
        document.add_paragraph(ws.title)
        # 根据Worksheet的行数和列数，在Word文件中创建合适大小的表格
        table = document.add_table(rows=len(rows), cols=len(rows[0]), style=choice(table_styles))
        # 从Worksheet读取数据，写入Word文件中的表格
        for irow, row in enumerate(rows):
            for icol, col in enumerate(row):
                table.cell(irow, icol).text = str(col.value)
    # 保存Word文件
    document.save(fn[:-4]+'docx')

# 调用函数，进行数据导入
xlsx2docx('测试文件.xlsx')

def docx2xlsx(fn):
    document = Document(fn)
    wb = Workbook()
    wb.remove_sheet(wb.worksheets[0])
    for index, table in enumerate(document.tables, start=1):
        ws = wb.create_sheet('sheet{}'.format(index))
        for row in table.rows:
            values = list(map(lambda cell:cell.text, row.cells))
            ws.append(values)
    wb.save(fn[:-5]+'_new.xlsx')
    
docx2xlsx('测试文件.docx')
